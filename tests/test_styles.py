from docx import Document
from docxcompose.composer import Composer
from utils import ComposedDocument
from utils import docx_path
from utils import FixtureDocument
import pytest


def test_contains_predefined_styles_in_masters_language(merged_styles):
    style_ids = [s.style_id for s in merged_styles.doc.styles]
    assert 'Heading1' in style_ids
    assert 'Heading1' in style_ids
    assert 'Strong' in style_ids
    assert 'Quote' in style_ids


def test_does_not_contain_predefined_styles_in_appended_language(merged_styles):
    style_ids = [s.style_id for s in merged_styles.doc.styles]
    assert 'berschrift1' not in style_ids
    assert 'berschrift2' not in style_ids
    assert 'Fett' not in style_ids
    assert 'Zitat' not in style_ids


def test_contains_custom_styles_from_both_docs(merged_styles):
    style_ids = [s.style_id for s in merged_styles.doc.styles]
    assert 'MyStyle1' in style_ids
    assert 'MyStyle1Char' in style_ids
    assert 'MeineFormatvorlage' in style_ids
    assert 'MeineFormatvorlageZchn' in style_ids


def test_contains_linked_styles(merged_styles):
    style_ids = [s.style_id for s in merged_styles.doc.styles]
    assert 'QuoteChar' in style_ids


def test_merged_styles_de():
    doc = FixtureDocument("styles_de.docx")
    composed = ComposedDocument(
        "styles_de.docx", "styles_en.docx")

    assert composed == doc


def test_merged_styles_en():
    doc = FixtureDocument("styles_en.docx")
    composed = ComposedDocument(
        "styles_en.docx", "styles_de.docx")

    assert composed == doc


def test_styles_are_not_switched_for_first_numbering_element():
    doc = FixtureDocument("switched_listing_style.docx")
    composed = ComposedDocument(
        "master_switched_listing_style.docx", "switched_listing_style.docx")

    assert composed == doc

def test_continue_when_no_styles():
    """Expects not to throw a type error"""
    doc = ComposedDocument("aatmay.docx", "aatmay.docx")


def test_preserve_document_styles_creates_unique_style_ids():
    """Test that preserve_document_styles creates unique style IDs for each document."""
    master_doc = Document(docx_path("styles_en.docx"))
    composer = Composer(master_doc, preserve_document_styles=True)
    
    # Get initial style count
    initial_style_count = len(master_doc.styles)
    
    # Append a document with potentially conflicting styles
    doc_to_append = Document(docx_path("styles_de.docx"))
    composer.append(doc_to_append)
    
    style_ids = [s.style_id for s in composer.doc.styles]
    
    # Check that new styles have unique IDs with doc counter prefix
    assert any('doc1_' in style_id for style_id in style_ids), \
        "Expected to find styles with 'doc1_' prefix"
    
    # Verify the number of styles increased
    assert len(composer.doc.styles) > initial_style_count, \
        "Expected more styles after appending with preserve_document_styles=True"


def test_preserve_document_styles_keeps_both_styles():
    """Test that both master and appended documents' styles are preserved."""
    master_doc = Document(docx_path("styles_en.docx"))
    composer = Composer(master_doc, preserve_document_styles=True)
    
    # Get master document's custom styles
    master_style_ids = [s.style_id for s in master_doc.styles]
    
    # Append a document
    doc_to_append = Document(docx_path("styles_de.docx"))
    appended_style_ids = [s.style_id for s in doc_to_append.styles]
    
    composer.append(doc_to_append)
    
    final_style_ids = [s.style_id for s in composer.doc.styles]
    
    # Check that master styles are still present
    assert 'MyStyle1' in final_style_ids, \
        "Master document custom style should be preserved"
    
    # Check that appended document's styles exist with unique IDs
    assert any('doc1_MeineFormatvorlage' in style_id for style_id in final_style_ids), \
        "Appended document custom style should be preserved with unique ID"


def test_preserve_document_styles_multiple_documents():
    """Test that multiple documents can be appended with unique style IDs."""
    master_doc = Document(docx_path("styles_en.docx"))
    composer = Composer(master_doc, preserve_document_styles=True)
    
    # Append first document
    composer.append(Document(docx_path("styles_de.docx")))
    
    # Append second document (same document again)
    composer.append(Document(docx_path("styles_de.docx")))
    
    style_ids = [s.style_id for s in composer.doc.styles]
    
    # Check for styles from both appends with different prefixes
    doc1_styles = [sid for sid in style_ids if 'doc1_' in sid]
    doc2_styles = [sid for sid in style_ids if 'doc2_' in sid]
    
    assert len(doc1_styles) > 0, "Expected styles from first appended document"
    assert len(doc2_styles) > 0, "Expected styles from second appended document"


def test_default_behavior_unchanged():
    """Test that default behavior (preserve_document_styles=False) remains unchanged."""
    master_doc = Document(docx_path("styles_en.docx"))
    composer = Composer(master_doc)  # Default: preserve_document_styles=False
    
    initial_style_ids = [s.style_id for s in master_doc.styles]
    
    doc_to_append = Document(docx_path("styles_de.docx"))
    composer.append(doc_to_append)
    
    final_style_ids = [s.style_id for s in composer.doc.styles]
    
    # Should not have any 'doc1_' prefixed styles in default mode
    assert not any('doc1_' in style_id for style_id in final_style_ids), \
        "Default behavior should not create prefixed style IDs"


@pytest.fixture
def merged_styles():
    composer = Composer(Document(docx_path("styles_en.docx")))
    composer.append(Document(docx_path("styles_de.docx")))
    return composer
